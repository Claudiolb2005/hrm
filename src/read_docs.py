"""Extrae texto de los .docx de dudas (párrafos + tablas)."""
import glob, os
import docx

for path in sorted(glob.glob("ref/hrm-docs/*.docx")):
    d = docx.Document(path)
    print(f"\n{'#'*74}\n# {os.path.basename(path)}\n{'#'*74}")
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            print(t)
    for ti, tbl in enumerate(d.tables):
        print(f"\n[TABLA {ti+1}]")
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                print("  | " + " | ".join(cells))
